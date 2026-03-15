"""DOCX text extraction using python-docx."""

from pathlib import Path

from docx import Document


def extract_text_from_docx(path: str | Path) -> str:
    """
    Extract plain text from a Word document (.docx).

    Args:
        path: Path to the DOCX file.

    Returns:
        Extracted text. Paragraphs separated by newlines.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file is not a valid DOCX.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    if path.suffix.lower() not in (".docx", ".doc"):
        raise ValueError(f"Expected .docx file, got: {path.suffix}")

    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    return "\n\n".join(paragraphs)


def extract_text_from_docx_with_meta(path: str | Path) -> tuple[str, dict]:
    """
    Extract text and metadata from a DOCX.

    Returns:
        Tuple of (full_text, metadata_dict). Metadata may include title, author, subject.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    doc = Document(path)
    text = extract_text_from_docx(path)

    core_props = doc.core_properties
    metadata = {
        "title": core_props.title,
        "author": core_props.author,
        "subject": core_props.subject,
        "keywords": core_props.keywords,
        "created": str(core_props.created) if core_props.created else None,
        "modified": str(core_props.modified) if core_props.modified else None,
    }

    return text, metadata
